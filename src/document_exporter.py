"""
Document Exporter - Export content to Word and HTML

This module exports generated content to Word documents and HTML files
suitable for direct use in content editors.

Features:
- Export to formatted Word documents (.docx)
- Export to clean HTML (without <html>, <head>, <body> tags)
- Preserve formatting (headings, bold, lists, etc.)
- Include SEO title and meta description separately
"""

import logging
from typing import Dict, List, Any
from pathlib import Path
import pandas as pd
import re
from html.parser import HTMLParser
from io import StringIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Word export will not be available.")

logger = logging.getLogger(__name__)


class LanguageDetector:
    """Detect text language and determine text direction."""
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Detect if text is primarily Persian/Farsi or English.
        
        Args:
            text: Text to analyze
            
        Returns:
            'persian' if Persian text, 'english' if English text
        """
        if not text or not text.strip():
            return 'english'  # Default to English
        
        # Count Persian characters (Arabic script)
        persian_chars = 0
        english_chars = 0
        
        for char in text:
            if '\u0600' <= char <= '\u06FF':  # Arabic/Persian Unicode range
                persian_chars += 1
            elif char.isalpha() and ord(char) < 128:  # Basic Latin
                english_chars += 1
        
        # If more Persian characters, it's Persian
        if persian_chars > english_chars:
            return 'persian'
        else:
            return 'english'
    
    @staticmethod
    def get_text_direction(text: str) -> str:
        """
        Get text direction based on language detection.
        
        Args:
            text: Text to analyze
            
        Returns:
            'rtl' for Persian text, 'ltr' for English text
        """
        language = LanguageDetector.detect_language(text)
        return 'rtl' if language == 'persian' else 'ltr'


class HTMLToWordConverter:
    """Convert HTML content to Word document format with language support."""
    
    def __init__(self, document: 'Document'):
        """
        Initialize converter.
        
        Args:
            document: python-docx Document instance
        """
        self.doc = document
        self.current_list = None
        self.language_detector = LanguageDetector()
    
    def parse_html_to_word(self, html_content: str):
        """
        Parse HTML content and add to Word document.
        
        Args:
            html_content: HTML content string
        """
        # Simple HTML parser for common tags
        # Remove HTML comments
        html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
        
        # Split by major tags
        parts = re.split(r'(<h[1-6][^>]*>.*?</h[1-6]>|<p[^>]*>.*?</p>|<ul[^>]*>.*?</ul>|<ol[^>]*>.*?</ol>)', 
                        html_content, flags=re.DOTALL)
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
            
            # Handle headings
            if re.match(r'<h[1-6]', part, re.IGNORECASE):
                self._add_heading(part)
            
            # Handle paragraphs
            elif re.match(r'<p', part, re.IGNORECASE):
                self._add_paragraph(part)
            
            # Handle unordered lists
            elif re.match(r'<ul', part, re.IGNORECASE):
                self._add_list(part, ordered=False)
            
            # Handle ordered lists
            elif re.match(r'<ol', part, re.IGNORECASE):
                self._add_list(part, ordered=True)
            
            # Plain text
            elif part and not part.startswith('<'):
                p = self.doc.add_paragraph(part)
                p.style = 'Normal'
    
    def _add_heading(self, html: str):
        """Add heading to document with proper text direction."""
        # Extract level
        level_match = re.search(r'<h([1-6])', html, re.IGNORECASE)
        level = int(level_match.group(1)) if level_match else 1
        
        # Extract text
        text = re.sub(r'<[^>]+>', '', html).strip()
        text = self._decode_html_entities(text)
        
        # Add to document
        heading = self.doc.add_heading(text, level=min(level, 3))
        
        # Set proper text direction based on language
        self._set_paragraph_direction(heading, text)
    
    def _add_paragraph(self, html: str):
        """Add paragraph to document with inline formatting and proper direction."""
        # Remove <p> tags
        text = re.sub(r'</?p[^>]*>', '', html, flags=re.IGNORECASE)
        
        # Create paragraph
        p = self.doc.add_paragraph()
        
        # Parse inline formatting
        self._parse_inline_formatting(text, p)
        
        # Set proper text direction based on language
        self._set_paragraph_direction(p, text)
    
    def _add_list(self, html: str, ordered: bool = False):
        """Add list to document with proper text direction."""
        # Extract list items
        items = re.findall(r'<li[^>]*>(.*?)</li>', html, re.DOTALL | re.IGNORECASE)
        
        for item in items:
            text = re.sub(r'<[^>]+>', '', item).strip()
            text = self._decode_html_entities(text)
            
            # Add as bullet point
            p = self.doc.add_paragraph(text, style='List Bullet' if not ordered else 'List Number')
            
            # Set proper text direction based on language
            self._set_paragraph_direction(p, text)
    
    def _parse_inline_formatting(self, html: str, paragraph):
        """Parse inline HTML formatting and add to paragraph with proper text direction."""
        # Split by tags
        parts = re.split(r'(<strong[^>]*>.*?</strong>|<b[^>]*>.*?</b>|<em[^>]*>.*?</em>|<i[^>]*>.*?</i>)', 
                        html, flags=re.DOTALL)
        
        for part in parts:
            if not part.strip():
                continue
            
            # Bold
            if re.match(r'<(strong|b)', part, re.IGNORECASE):
                text = re.sub(r'<[^>]+>', '', part).strip()
                text = self._decode_html_entities(text)
                run = paragraph.add_run(text)
                run.bold = True
                # Set run direction based on text content
                self._set_run_direction(run, text)
            
            # Italic
            elif re.match(r'<(em|i)', part, re.IGNORECASE):
                text = re.sub(r'<[^>]+>', '', part).strip()
                text = self._decode_html_entities(text)
                run = paragraph.add_run(text)
                run.italic = True
                # Set run direction based on text content
                self._set_run_direction(run, text)
            
            # Plain text
            else:
                text = re.sub(r'<[^>]+>', '', part).strip()
                text = self._decode_html_entities(text)
                if text:
                    run = paragraph.add_run(text)
                    # Set run direction based on text content
                    self._set_run_direction(run, text)
    
    def _set_paragraph_direction(self, paragraph, text: str):
        """
        Set paragraph text direction based on language detection.
        
        Args:
            paragraph: Word paragraph object
            text: Text content to analyze
        """
        try:
            # Detect language and set direction
            direction = self.language_detector.get_text_direction(text)
            
            if direction == 'rtl':
                # Set right-to-left alignment for Persian text
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # Set RTL direction in paragraph properties
                self._set_rtl_direction(paragraph)
            else:
                # Set left-to-right alignment for English text
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Set LTR direction in paragraph properties
                self._set_ltr_direction(paragraph)
                
        except Exception as e:
            logger.warning(f"Failed to set paragraph direction: {e}")
            # Default to left alignment
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _set_rtl_direction(self, paragraph):
        """Set right-to-left text direction for paragraph."""
        try:
            # Get paragraph element
            p_element = paragraph._element
            
            # Set bidi (bidirectional) property
            bidi = OxmlElement('w:bidi')
            p_element.get_or_add_pPr().append(bidi)
            
            # Set text direction
            text_direction = OxmlElement('w:textDirection')
            text_direction.set(qn('w:val'), 'rl')
            p_element.get_or_add_pPr().append(text_direction)
            
        except Exception as e:
            logger.warning(f"Failed to set RTL direction: {e}")
    
    def _set_ltr_direction(self, paragraph):
        """Set left-to-right text direction for paragraph."""
        try:
            # Get paragraph element
            p_element = paragraph._element
            
            # Set text direction to LTR
            text_direction = OxmlElement('w:textDirection')
            text_direction.set(qn('w:val'), 'lr')
            p_element.get_or_add_pPr().append(text_direction)
            
        except Exception as e:
            logger.warning(f"Failed to set LTR direction: {e}")
    
    def _set_run_direction(self, run, text: str):
        """
        Set text direction for a run based on language detection.
        
        Args:
            run: Word run object
            text: Text content to analyze
        """
        try:
            # Detect language and set direction
            direction = self.language_detector.get_text_direction(text)
            
            if direction == 'rtl':
                # Set RTL direction for Persian text
                self._set_run_rtl_direction(run)
            else:
                # Set LTR direction for English text
                self._set_run_ltr_direction(run)
                
        except Exception as e:
            logger.warning(f"Failed to set run direction: {e}")
    
    def _set_run_rtl_direction(self, run):
        """Set right-to-left text direction for run."""
        try:
            # Get run element
            r_element = run._element
            
            # Set bidi (bidirectional) property
            bidi = OxmlElement('w:bidi')
            r_element.get_or_add_rPr().append(bidi)
            
            # Set text direction
            text_direction = OxmlElement('w:textDirection')
            text_direction.set(qn('w:val'), 'rl')
            r_element.get_or_add_rPr().append(text_direction)
            
        except Exception as e:
            logger.warning(f"Failed to set run RTL direction: {e}")
    
    def _set_run_ltr_direction(self, run):
        """Set left-to-right text direction for run."""
        try:
            # Get run element
            r_element = run._element
            
            # Set text direction to LTR
            text_direction = OxmlElement('w:textDirection')
            text_direction.set(qn('w:val'), 'lr')
            r_element.get_or_add_rPr().append(text_direction)
            
        except Exception as e:
            logger.warning(f"Failed to set run LTR direction: {e}")
    
    def _decode_html_entities(self, text: str) -> str:
        """Decode HTML entities."""
        import html
        return html.unescape(text)


class DocumentExporter:
    """Export generated content to Word and HTML formats."""
    
    def __init__(self, output_dir: str = "output/documents"):
        """
        Initialize Document Exporter.
        
        Args:
            output_dir: Output directory for documents
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        if not DOCX_AVAILABLE:
            logger.warning("⚠️  python-docx not available. Word export disabled.")
        
        logger.info("✅ Document Exporter initialized")
    
    def export_content_to_word(
        self,
        title: str,
        meta_description: str,
        content_html: str,
        output_filename: str
    ) -> str:
        """
        Export content to Word document.
        
        Args:
            title: SEO title
            meta_description: Meta description
            content_html: HTML content
            output_filename: Output filename (without extension)
            
        Returns:
            Path to created Word file
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed. Cannot create Word document.")
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")
        
        try:
            # Create document
            doc = Document()
            
            # Add title section
            doc.add_heading('SEO Information', level=1)
            
            # Title
            p = doc.add_paragraph()
            p.add_run('Title: ').bold = True
            p.add_run(title)
            
            # Meta Description
            p = doc.add_paragraph()
            p.add_run('Meta Description: ').bold = True
            p.add_run(meta_description)
            
            # Separator
            doc.add_paragraph('_' * 60)
            
            # Add main content
            doc.add_heading('Content', level=1)
            
            # Convert HTML to Word
            converter = HTMLToWordConverter(doc)
            converter.parse_html_to_word(content_html)
            
            # Save
            output_path = self.output_dir / f"{output_filename}.docx"
            doc.save(str(output_path))
            
            logger.info(f"✅ Created Word document: {output_path.name}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            raise
    
    def export_content_to_html(
        self,
        title: str,
        meta_description: str,
        content_html: str,
        output_filename: str
    ) -> str:
        """
        Export content to HTML file (editor-ready, no <html> wrapper).
        
        Args:
            title: SEO title
            meta_description: Meta description
            content_html: HTML content
            output_filename: Output filename (without extension)
            
        Returns:
            Path to created HTML file
        """
        try:
            # Clean HTML content (remove document-level tags)
            clean_content = self._clean_html_for_editor(content_html)
            
            # Build editor-ready HTML (without comments, clean content only)
            html_output = clean_content
            
            # Save
            output_path = self.output_dir / f"{output_filename}.html"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_output)
            
            logger.info(f"✅ Created HTML file: {output_path.name}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to create HTML file: {e}")
            raise
    
    def _clean_html_for_editor(self, html_content: str) -> str:
        """
        Clean HTML for editor use (remove document tags).
        
        Args:
            html_content: Raw HTML
            
        Returns:
            Cleaned HTML
        """
        # Remove document-level tags
        html_content = re.sub(r'<!DOCTYPE[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<html[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</html>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<head[^>]*>.*?</head>', '', html_content, flags=re.IGNORECASE | re.DOTALL)
        html_content = re.sub(r'<body[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</body>', '', html_content, flags=re.IGNORECASE)
        
        # Clean whitespace
        html_content = re.sub(r'\n\s*\n', '\n\n', html_content)
        html_content = html_content.strip()
        
        return html_content
    
    def export_batch_from_dataframe(
        self,
        df: pd.DataFrame,
        base_filename: str = "content"
    ) -> Dict[str, List[str]]:
        """
        Export all content from DataFrame to Word and HTML files.
        
        Args:
            df: DataFrame with columns: SEO_Title, Meta_Description, Generated_Content
            base_filename: Base filename for outputs
            
        Returns:
            Dictionary with lists of created files
        """
        results = {
            'word_files': [],
            'html_files': [],
            'errors': []
        }
        
        print(f"\n{'='*70}")
        print(f"📄 Exporting Content to Word & HTML")
        print(f"{'='*70}\n")
        
        from tqdm import tqdm
        
        for idx, row in tqdm(df.iterrows(), total=len(df), desc="Exporting files"):
            try:
                title = row.get('SEO_Title', f'Content {idx+1}')
                meta_desc = row.get('Meta_Description', '')
                content = row.get('Generated_Content', '')
                
                if not content or pd.isna(content):
                    logger.warning(f"   Row {idx}: No content to export, skipping")
                    continue
                
                # Clean filename
                safe_title = re.sub(r'[^\w\s-]', '', str(title))
                safe_title = re.sub(r'[-\s]+', '-', safe_title)[:50]
                filename = f"{base_filename}_{idx+1}_{safe_title}"
                
                # Export to Word
                if DOCX_AVAILABLE:
                    try:
                        word_file = self.export_content_to_word(
                            title=str(title),
                            meta_description=str(meta_desc),
                            content_html=str(content),
                            output_filename=filename
                        )
                        results['word_files'].append(word_file)
                    except Exception as e:
                        logger.error(f"   Word export failed for row {idx}: {e}")
                        results['errors'].append(f"Row {idx} Word: {str(e)}")
                
                # Export to HTML
                try:
                    html_file = self.export_content_to_html(
                        title=str(title),
                        meta_description=str(meta_desc),
                        content_html=str(content),
                        output_filename=filename
                    )
                    results['html_files'].append(html_file)
                except Exception as e:
                    logger.error(f"   HTML export failed for row {idx}: {e}")
                    results['errors'].append(f"Row {idx} HTML: {str(e)}")
                
            except Exception as e:
                logger.error(f"   Row {idx} export failed: {e}")
                results['errors'].append(f"Row {idx}: {str(e)}")
        
        # Summary
        print(f"\n{'='*70}")
        print(f"✅ Export Complete!")
        print(f"{'='*70}")
        if DOCX_AVAILABLE:
            print(f"   📝 Word files: {len(results['word_files'])}")
        print(f"   🌐 HTML files: {len(results['html_files'])}")
        if results['errors']:
            print(f"   ⚠️  Errors: {len(results['errors'])}")
        print(f"   📁 Output directory: {self.output_dir.absolute()}")
        print(f"{'='*70}\n")
        
        return results

